"""Render an artifact's markdown (from `views._RENDERERS`) into a downloadable .docx or .pdf.

The artifact markdown is intentionally simple — a bold section heading, bullet lists (some
indented one level), and short paragraphs — so a tiny block parser is enough. Both exporters share
that parser so DOCX and PDF stay consistent.
"""

import re
from io import BytesIO

# ── Markdown → blocks ─────────────────────────────────────────────────

_BOLD_ONLY = re.compile(r"^\*\*(?P<text>.+)\*\*$")
_INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")


def _parse_blocks(markdown: str) -> list[tuple[str, str]]:
    """Return (kind, text) blocks: kind in {heading, bullet, subbullet, para}."""
    blocks: list[tuple[str, str]] = []
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("    - ") or line.startswith("\t- "):
            blocks.append(("subbullet", line.strip()[2:].strip()))
        elif line.lstrip().startswith("- "):
            blocks.append(("bullet", line.lstrip()[2:].strip()))
        elif _BOLD_ONLY.match(line.strip()):
            blocks.append(("heading", _BOLD_ONLY.match(line.strip()).group("text")))
        else:
            blocks.append(("para", line.strip()))
    return blocks


def _strip_leading_heading(blocks: list[tuple[str, str]]) -> list[tuple[str, str]]:
    """Drop the first block if it's just the section name (we render `title` as the H1 instead)."""
    if blocks and blocks[0][0] == "heading":
        return blocks[1:]
    return blocks


# ── Inline formatting helpers ─────────────────────────────────────────


def _inline_runs(text: str) -> list[tuple[str, bool, bool]]:
    """Split text into (chunk, bold, italic) runs for python-docx."""
    runs: list[tuple[str, bool, bool]] = []
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            runs.append((text[pos : m.start()], False, False))
        if m.group(1) is not None:
            runs.append((m.group(1), True, False))
        else:
            runs.append((m.group(2), False, True))
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], False, False))
    return runs or [(text, False, False)]


def _inline_html(text: str) -> str:
    """Escape XML then convert **bold**/*italic* to reportlab mini-HTML tags."""
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
    safe = re.sub(r"\*(.+?)\*", r"<i>\1</i>", safe)
    return safe


# ── DOCX ──────────────────────────────────────────────────────────────


def markdown_to_docx(title: str, markdown: str) -> bytes:
    import docx
    from docx.shared import Pt, RGBColor

    document = docx.Document()
    heading = document.add_heading(title, level=0)
    heading.runs[0].font.color.rgb = RGBColor(0x00, 0x00, 0x75)  # navy brand

    for kind, text in _strip_leading_heading(_parse_blocks(markdown)):
        if kind == "heading":
            document.add_heading(text, level=1)
        elif kind == "bullet":
            para = document.add_paragraph(style="List Bullet")
            _add_runs(para, text)
        elif kind == "subbullet":
            para = document.add_paragraph(style="List Bullet 2")
            _add_runs(para, text)
        else:
            para = document.add_paragraph()
            _add_runs(para, text)
        if kind == "para":
            para.paragraph_format.space_after = Pt(6)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _add_runs(paragraph, text: str) -> None:
    for chunk, bold, italic in _inline_runs(text):
        run = paragraph.add_run(chunk)
        run.bold = bold
        run.italic = italic


# ── PDF ───────────────────────────────────────────────────────────────


def markdown_to_pdf(title: str, markdown: str) -> bytes:
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

    styles = getSampleStyleSheet()
    navy = HexColor("#000075")
    title_style = ParagraphStyle(
        "DocTitle", parent=styles["Title"], textColor=navy, alignment=TA_LEFT, fontSize=20, spaceAfter=10
    )
    h1 = ParagraphStyle("H1", parent=styles["Heading2"], textColor=navy, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontSize=10.5, leading=15)

    flow = [Paragraph(_inline_html(title), title_style), Spacer(1, 4)]
    pending_bullets: list = []

    def flush_bullets() -> None:
        nonlocal pending_bullets
        if pending_bullets:
            flow.append(ListFlowable(pending_bullets, bulletType="bullet", leftIndent=14))
            pending_bullets = []

    for kind, text in _strip_leading_heading(_parse_blocks(markdown)):
        if kind in ("bullet", "subbullet"):
            indent = 10 if kind == "subbullet" else 0
            pending_bullets.append(
                ListItem(Paragraph(_inline_html(text), body), leftIndent=indent)
            )
        else:
            flush_bullets()
            flow.append(Paragraph(_inline_html(text), h1 if kind == "heading" else body))
    flush_bullets()

    buffer = BytesIO()
    SimpleDocTemplate(
        buffer, pagesize=A4, leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=18 * mm
    ).build(flow)
    return buffer.getvalue()


def render(fmt: str, title: str, markdown: str) -> tuple[bytes, str]:
    """Render and return (bytes, mime_type) for fmt in {'docx', 'pdf'}."""
    if fmt == "docx":
        return (
            markdown_to_docx(title, markdown),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if fmt == "pdf":
        return markdown_to_pdf(title, markdown), "application/pdf"
    raise ValueError(f"Unsupported format: {fmt!r}")
